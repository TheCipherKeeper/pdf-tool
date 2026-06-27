"""MS Word (.docx) operations via python-docx + LibreOffice headless."""
from __future__ import annotations

import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from ..backends import require


def docx_info(path: Path) -> dict:
    """Return a summary dict: paragraphs, words, images, core properties."""
    doc = Document(str(path))
    paragraphs = len(doc.paragraphs)
    words = sum(len(p.text.split()) for p in doc.paragraphs if p.text.strip())
    # Inline images count (rough: count drawing/inline elements)
    images = 0
    for el in doc.element.iter(qn("w:drawing")):
        images += 1
    core = doc.core_properties
    return {
        "paragraphs": paragraphs,
        "words": words,
        "images": images,
        "title": core.title or "",
        "author": core.author or "",
        "subject": core.subject or "",
        "created": str(core.created) if core.created else "",
        "modified": str(core.modified) if core.modified else "",
    }


def docx_text(path: Path) -> str:
    """Extract all text, joining paragraph runs and preserving paragraph breaks.

    Walks the body in document order (paragraphs + tables) rather than only
    `doc.paragraphs`, so table cell text is included.
    """
    doc = Document(str(path))
    out: list[str] = []
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag
        if tag == qn("w:p"):
            # Reconstruct text from all w:t runs in this paragraph
            text = "".join(t.text or "" for t in child.iter(qn("w:t")))
            out.append(text)
        elif tag == qn("w:tbl"):
            # Flatten table cell text, tab-separated per cell, newline per row
            for row in child.iter(qn("w:tr")):
                cells = []
                for cell in row.iter(qn("w:tc")):
                    cell_text = "".join(t.text or "" for t in cell.iter(qn("w:t")))
                    cells.append(cell_text)
                out.append("\t".join(cells))
    return "\n".join(out)


def docx_merge(files: list[Path], out: Path) -> int:
    """Concatenate docx files by appending body content with a page break between.

    Returns the number of files merged. This is a structural merge (preserves
    text + runs); complex layout/headers may not transfer perfectly.
    """
    base = Document(str(files[0]))
    for extra in files[1:]:
        add_page_break(base)
        other = Document(str(extra))
        for element in other.element.body:
            # Skip the final sectPr (section properties) of the appended doc
            if element.tag == qn("w:sectPr"):
                continue
            base.element.body.append(element)
    out.parent.mkdir(parents=True, exist_ok=True)
    base.save(str(out))
    return len(files)


def add_page_break(doc: Document) -> None:
    """Append an explicit page break paragraph to a document."""
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def libreoffice_convert(src: Path, target_ext: str, out: Path) -> Path:
    """Convert a document via LibreOffice headless.

    `target_ext` is the LibreOffice output filter name (e.g. 'pdf', 'docx',
    'docx:"MS Word 2007 XML"'). LibreOffice writes to an output directory using
    the input stem + new extension; we then move it to the requested `out`.
    """
    soffice = require("soffice")
    with tempfile.TemporaryDirectory(prefix="loconv_") as tmpdir:
        args = [
            str(soffice),
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            f"--convert-to",
            target_ext,
            "--outdir",
            tmpdir,
            str(src),
        ]
        result = subprocess.run(args, capture_output=True, text=True)
        if result.returncode != 0:
            raise RuntimeError(
                f"LibreOffice conversion failed (exit {result.returncode}):\n"
                f"{result.stderr or result.stdout}"
            )
        produced = Path(tmpdir) / (src.stem + "." + target_ext.split(":")[0].strip('"'))
        if not produced.exists():
            # Fallback: pick the first non-source file produced
            produced_files = [p for p in Path(tmpdir).iterdir() if p != src]
            if not produced_files:
                raise RuntimeError(
                    "LibreOffice reported success but produced no output file.\n"
                    f"stdout: {result.stdout}"
                )
            produced = produced_files[0]
        out.parent.mkdir(parents=True, exist_ok=True)
        if out.exists():
            out.unlink()
        shutil.move(str(produced), str(out))
    return out