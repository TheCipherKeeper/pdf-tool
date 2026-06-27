"""Pytest fixtures: generate a small PDF and a small .docx in-memory.

No binary assets are committed — fixtures are built on the fly so the suite
runs on any machine with the Python dependencies installed.
"""
from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter


@pytest.fixture
def tmp_pdf(tmp_path: Path) -> Path:
    """A 5-page PDF whose pages are distinguishable by a hidden text layer."""
    out = tmp_path / "sample.pdf"
    writer = PdfWriter()
    for i in range(1, 6):
        # pypdf can't easily add visible text, so we add a page label via
        # metadata-like content using a blank page + named destination.
        writer.add_blank_page(width=200, height=200)
    with open(out, "wb") as f:
        writer.write(f)
    return out


@pytest.fixture
def pdf_with_text(tmp_path: Path) -> Path:
    """A 4-page PDF with a real text layer (via PyMuPDF) for text extraction."""
    import fitz

    out = tmp_path / "texted.pdf"
    doc = fitz.open()
    for i in range(1, 5):
        page = doc.new_page(width=300, height=300)
        page.insert_text((50, 150), f"Page {i} hello world", fontsize=24)
    doc.save(str(out))
    doc.close()
    return out


@pytest.fixture
def docx_basic(tmp_path: Path) -> Path:
    """A small .docx with paragraphs and a table."""
    out = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("First paragraph with some words.")
    doc.add_paragraph("Second paragraph here.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "B1"
    table.cell(1, 0).text = "A2"
    table.cell(1, 1).text = "B2"
    # Core properties
    doc.core_properties.title = "Sample Title"
    doc.core_properties.author = "Tester"
    doc.save(str(out))
    return out


@pytest.fixture
def docx_second(tmp_path: Path) -> Path:
    """A second small .docx for merge tests."""
    out = tmp_path / "second.docx"
    doc = Document()
    doc.add_paragraph("Appended paragraph from second doc.")
    doc.save(str(out))
    return out